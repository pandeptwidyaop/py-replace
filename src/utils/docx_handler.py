"""
Module untuk handling file DOCX - load, scan, dan save
"""
from docx import Document
from docx.text.paragraph import Paragraph
from docx.table import Table, _Cell
from docx.shared import Inches
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from typing import Set, Dict, List, Tuple
from .placeholder import PlaceholderHandler
from .image_handler import ImageHandler
import re


class DocxHandler:
    """Handler untuk operasi file DOCX"""

    def __init__(self, file_path: str = None):
        """
        Inisialisasi DocxHandler

        Args:
            file_path: Path ke file DOCX yang akan diload
        """
        self.file_path = file_path
        self.document = None
        if file_path:
            self.load(file_path)

    def load(self, file_path: str):
        """
        Load dokumen DOCX

        Args:
            file_path: Path ke file DOCX
        """
        self.file_path = file_path
        self.document = Document(file_path)

    def find_all_placeholders(self) -> Set[str]:
        """
        Menemukan semua TEXT placeholder dalam dokumen (${})

        Returns:
            Set dari nama placeholder yang ditemukan
        """
        if not self.document:
            return set()

        placeholders = set()

        # Scan paragraphs
        for paragraph in self.document.paragraphs:
            placeholders.update(
                PlaceholderHandler.find_placeholders(paragraph.text)
            )

        # Scan tables
        for table in self.document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        placeholders.update(
                            PlaceholderHandler.find_placeholders(paragraph.text)
                        )

        # Scan headers
        for section in self.document.sections:
            header = section.header
            for paragraph in header.paragraphs:
                placeholders.update(
                    PlaceholderHandler.find_placeholders(paragraph.text)
                )

            # Scan footers
            footer = section.footer
            for paragraph in footer.paragraphs:
                placeholders.update(
                    PlaceholderHandler.find_placeholders(paragraph.text)
                )

        return placeholders

    def find_all_image_placeholders(self) -> Set[str]:
        """
        Menemukan semua IMAGE placeholder dalam dokumen (@{})

        Returns:
            Set dari nama image placeholder yang ditemukan
        """
        if not self.document:
            return set()

        placeholders = set()

        # Scan paragraphs
        for paragraph in self.document.paragraphs:
            placeholders.update(
                PlaceholderHandler.find_image_placeholders(paragraph.text)
            )

        # Scan tables
        for table in self.document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        placeholders.update(
                            PlaceholderHandler.find_image_placeholders(paragraph.text)
                        )

        # Scan headers and footers
        for section in self.document.sections:
            # Headers
            for paragraph in section.header.paragraphs:
                placeholders.update(
                    PlaceholderHandler.find_image_placeholders(paragraph.text)
                )

            # Footers
            for paragraph in section.footer.paragraphs:
                placeholders.update(
                    PlaceholderHandler.find_image_placeholders(paragraph.text)
                )

        return placeholders

    def find_all_placeholders_with_types(self) -> Tuple[Set[str], Set[str]]:
        """
        Menemukan semua placeholder (text dan image) dalam dokumen

        Returns:
            Tuple (text_placeholders, image_placeholders)
        """
        text_placeholders = self.find_all_placeholders()
        image_placeholders = self.find_all_image_placeholders()
        return text_placeholders, image_placeholders

    def replace_placeholders(self, replacements: Dict[str, str]):
        """
        Mengganti semua placeholder dalam dokumen

        Args:
            replacements: Dictionary mapping placeholder -> nilai pengganti
        """
        if not self.document:
            return

        # Replace in paragraphs
        for paragraph in self.document.paragraphs:
            self._replace_in_paragraph(paragraph, replacements)

        # Replace in tables
        for table in self.document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self._replace_in_paragraph(paragraph, replacements)

        # Replace in headers and footers
        for section in self.document.sections:
            # Headers
            for paragraph in section.header.paragraphs:
                self._replace_in_paragraph(paragraph, replacements)

            # Footers
            for paragraph in section.footer.paragraphs:
                self._replace_in_paragraph(paragraph, replacements)

    def _replace_in_paragraph(self, paragraph: Paragraph, replacements: Dict[str, str]):
        """
        Mengganti placeholder dalam satu paragraph dengan mempertahankan formatting

        Args:
            paragraph: Paragraph object dari python-docx
            replacements: Dictionary mapping placeholder -> nilai pengganti
        """
        # Build full text to detect placeholders that may span multiple runs
        full_text = ''.join(run.text for run in paragraph.runs)

        # Check if there are any placeholders
        if not PlaceholderHandler.find_placeholders(full_text):
            return  # No placeholders, nothing to replace

        # Find all placeholders and their positions
        placeholder_matches = []
        for placeholder, value in replacements.items():
            pattern = r'\$\{' + re.escape(placeholder) + r'\}'
            for match in re.finditer(pattern, full_text):
                placeholder_matches.append({
                    'start': match.start(),
                    'end': match.end(),
                    'placeholder': placeholder,
                    'value': value,
                    'original': match.group()
                })

        if not placeholder_matches:
            return  # No matches to replace

        # Sort by position
        placeholder_matches.sort(key=lambda x: x['start'])

        # Build new runs with preserved formatting
        new_runs_data = []
        current_pos = 0
        run_position = 0

        for run in paragraph.runs:
            run_start = run_position
            run_end = run_position + len(run.text)
            run_text = run.text

            # Check if this run contains any part of placeholders
            run_segments = []
            segment_start = 0

            for match in placeholder_matches:
                match_start = match['start']
                match_end = match['end']

                # Check if placeholder overlaps with this run
                if match_start < run_end and match_end > run_start:
                    # Calculate the portion of placeholder in this run
                    local_match_start = max(0, match_start - run_start)
                    local_match_end = min(len(run_text), match_end - run_start)

                    # Add text before placeholder (if any)
                    if segment_start < local_match_start:
                        run_segments.append({
                            'text': run_text[segment_start:local_match_start],
                            'is_replacement': False
                        })

                    # Check if this is the run where placeholder starts
                    if match_start >= run_start and match_start < run_end:
                        # This run contains the start of the placeholder, do replacement here
                        run_segments.append({
                            'text': match['value'],
                            'is_replacement': True
                        })

                    segment_start = local_match_end

            # Add remaining text in run
            if segment_start < len(run_text):
                run_segments.append({
                    'text': run_text[segment_start:],
                    'is_replacement': False
                })

            # Store run data with formatting
            if run_segments:
                for segment in run_segments:
                    if segment['text']:  # Only add non-empty segments
                        new_runs_data.append({
                            'text': segment['text'],
                            'bold': run.bold,
                            'italic': run.italic,
                            'underline': run.underline,
                            'font_name': run.font.name,
                            'font_size': run.font.size,
                            'font_color': run.font.color.rgb if run.font.color.rgb else None,
                        })

            run_position = run_end

        # Clear all existing runs
        for i in range(len(paragraph.runs) - 1, -1, -1):
            paragraph._element.remove(paragraph.runs[i]._element)

        # Add new runs with preserved formatting
        for run_data in new_runs_data:
            new_run = paragraph.add_run(run_data['text'])

            # Apply formatting
            if run_data['bold'] is not None:
                new_run.bold = run_data['bold']
            if run_data['italic'] is not None:
                new_run.italic = run_data['italic']
            if run_data['underline'] is not None:
                new_run.underline = run_data['underline']
            if run_data['font_name']:
                new_run.font.name = run_data['font_name']
            if run_data['font_size']:
                new_run.font.size = run_data['font_size']
            if run_data['font_color']:
                new_run.font.color.rgb = run_data['font_color']

    def save(self, output_path: str):
        """
        Simpan dokumen ke file

        Args:
            output_path: Path output file
        """
        if self.document:
            self.document.save(output_path)

    def replace_image_placeholders(self, image_replacements: Dict[str, str],
                                   width_inches: float = 3.0) -> Tuple[int, List[str]]:
        """
        Mengganti image placeholder dengan actual images

        Args:
            image_replacements: Dictionary mapping placeholder -> image path/URL
            width_inches: Lebar default image dalam inches

        Returns:
            Tuple (success_count, error_messages)
        """
        if not self.document:
            return 0, ["Document not loaded"]

        success_count = 0
        errors = []
        temp_files = []

        try:
            for placeholder, image_path in image_replacements.items():
                # Get and validate image path
                final_path, is_temp, error = ImageHandler.get_image_path(image_path)

                if error:
                    errors.append(f"{placeholder}: {error}")
                    continue

                if is_temp:
                    temp_files.append(final_path)

                # Replace in paragraphs
                replaced = self._replace_image_in_paragraphs(
                    self.document.paragraphs,
                    placeholder,
                    final_path,
                    width_inches
                )

                # Replace in tables
                for table in self.document.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            replaced += self._replace_image_in_paragraphs(
                                cell.paragraphs,
                                placeholder,
                                final_path,
                                width_inches
                            )

                # Replace in headers and footers
                for section in self.document.sections:
                    # Headers
                    replaced += self._replace_image_in_paragraphs(
                        section.header.paragraphs,
                        placeholder,
                        final_path,
                        width_inches
                    )

                    # Footers
                    replaced += self._replace_image_in_paragraphs(
                        section.footer.paragraphs,
                        placeholder,
                        final_path,
                        width_inches
                    )

                if replaced > 0:
                    success_count += 1

        finally:
            # Cleanup temp files
            for temp_file in temp_files:
                ImageHandler.cleanup_temp_file(temp_file)

        return success_count, errors

    def _replace_image_in_paragraphs(self, paragraphs, placeholder: str,
                                     image_path: str, width_inches: float) -> int:
        """
        Replace image placeholder dalam list of paragraphs

        Args:
            paragraphs: List of paragraphs
            placeholder: Nama placeholder
            image_path: Path ke image file
            width_inches: Lebar image

        Returns:
            Number of replacements made
        """
        replacements = 0
        pattern = r'@\{' + re.escape(placeholder) + r'\}'

        for paragraph in paragraphs:
            if re.search(pattern, paragraph.text):
                # Clear paragraph text
                paragraph.text = ""

                # Add image
                try:
                    run = paragraph.add_run()
                    run.add_picture(image_path, width=Inches(width_inches))
                    replacements += 1
                except Exception as e:
                    # If failed, restore placeholder with error note
                    paragraph.text = f"@{{{placeholder}}} [Error: {str(e)}]"

        return replacements

    def get_document_info(self) -> Dict[str, any]:
        """
        Mendapatkan informasi dokumen

        Returns:
            Dictionary berisi informasi dokumen
        """
        if not self.document:
            return {}

        return {
            'paragraphs': len(self.document.paragraphs),
            'tables': len(self.document.tables),
            'sections': len(self.document.sections),
        }
